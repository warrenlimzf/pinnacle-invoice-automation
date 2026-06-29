"""Owns the per-bank verification Word doc (banks/<bank>/<bank>_verification.docx).

For each statement it appends: the figures we read, the LGT add-back breakdown
(if any), and the screenshot(s) of the exact PDF section — so your colleague opens
this next to the Excel and eyeballs that everything matches.
"""
from datetime import datetime
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from shared.model import ClientResult, FieldHit

_GREY = RGBColor(0x55, 0x55, 0x55)


def _doc(docx_path: Path) -> Document:
    if docx_path.exists():
        return Document(str(docx_path))
    doc = Document()
    doc.add_heading("Verification snapshots", level=0)
    p = doc.add_paragraph(
        "Each entry shows the exact section of the statement where the NAV figures "
        "were read. Compare against the Excel tab for this bank."
    )
    p.runs[0].italic = True
    return doc


def append_verification(docx_path, result: ClientResult,
                        images: List[Tuple[FieldHit, Path]], pdf_path) -> None:
    docx_path, pdf_path = Path(docx_path), Path(pdf_path)
    doc = _doc(docx_path)

    doc.add_heading(pdf_path.stem, level=1)

    meta = doc.add_paragraph()
    meta.add_run("Account No: ").bold = True
    meta.add_run(result.account_no or "(to be filled)")
    meta.add_run(f"     Processed: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    vals = doc.add_paragraph()
    vals.add_run("Gross NAV: ").bold = True
    vals.add_run(f"{_fmt(result.gross_nav)}     ")
    vals.add_run("Net NAV: ").bold = True
    vals.add_run(f"{_fmt(result.net_nav)}")

    if result.addbacks:
        doc.add_paragraph("Add-backs used to derive Gross NAV (LGT):").runs[0].bold = True
        for ab in result.addbacks:
            doc.add_paragraph(f"   {ab.label}:  {_fmt(ab.value)}", style=None)

    if result.flags:
        f = doc.add_paragraph()
        run = f.add_run("⚠ " + "; ".join(result.flags))
        run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
        run.font.size = Pt(10)

    for hit, png in images:
        cap = doc.add_paragraph()
        run = cap.add_run(f"{hit.label}  (page {hit.page + 1})")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = _GREY
        if png and Path(png).exists():
            doc.add_picture(str(png), width=Inches(6.0))

    doc.add_paragraph("")
    doc.save(str(docx_path))


def _fmt(v):
    return f"{v:,.2f}" if isinstance(v, (int, float)) else "(not found)"
